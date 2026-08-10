"""
Membuat dokumen notulensi (.docx) mengikuti format resmi
'template-notule-kegiatan.md' yang disediakan BPS Kabupaten Sanggau:
kop surat berulang di setiap halaman, tabel info rapat, tabel peserta,
narasi pendahuluan, pembahasan & keputusan, tabel tindak lanjut, blok
tanda tangan (Kepala & Notulis), serta lampiran foto dokumentasi kegiatan.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config import settings
from ..utils.indo_date import (
    format_tanggal_lengkap, format_tanggal_singkat, format_jam, hitung_durasi,
)


def _add_letterhead(doc: Document):
    """Kop surat, dipasang sebagai header halaman agar berulang otomatis
    di setiap halaman (persis seperti pada template asli)."""
    header = doc.sections[0].header
    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(settings.NAMA_INSTANSI)
    run1.bold = True
    run1.italic = True
    run1.font.size = Pt(12)

    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(settings.ALAMAT_INSTANSI)
    run2.font.size = Pt(8)

    # garis pemisah tipis di bawah kop
    p3 = header.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    run3 = p3.add_run("_" * 95)
    run3.font.size = Pt(6)


def _info_table(doc: Document, meeting):
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells = table.rows[0].cells
    cells[0].text = "Unit Kerja"
    cells[1].text = meeting.unit_kerja or settings.UNIT_KERJA_DEFAULT
    cells[2].text = "Tanggal"
    cells[3].text = format_tanggal_singkat(meeting.tanggal)

    cells = table.rows[1].cells
    cells[0].text = "Topik"
    cells[1].text = meeting.judul_rapat
    cells[2].text = "Tempat"
    cells[3].text = meeting.lokasi or "-"

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].runs[0].bold = True
    return table


def _peserta_table(doc: Document, peserta_list):
    doc.add_paragraph("")
    heading = doc.add_paragraph()
    run = heading.add_run("Peserta Rapat")
    run.bold = True

    rows_needed = (len(peserta_list) + 1) // 2 if peserta_list else 1
    table = doc.add_table(rows=rows_needed + 1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Nama", "Jabatan", "Nama", "Jabatan"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True

    for i in range(rows_needed):
        row = table.rows[i + 1].cells
        left = peserta_list[i * 2] if i * 2 < len(peserta_list) else None
        right = peserta_list[i * 2 + 1] if i * 2 + 1 < len(peserta_list) else None
        row[0].text = left.nama if left else ""
        row[1].text = left.jabatan if left else ""
        row[2].text = right.nama if right else ""
        row[3].text = right.jabatan if right else ""
    return table


def _pendahuluan_paragraph(doc: Document, meeting, pimpinan):
    doc.add_heading("Pendahuluan", level=2)
    tanggal_lengkap = format_tanggal_lengkap(meeting.tanggal)
    jam_mulai = format_jam(meeting.waktu_mulai)
    jam_selesai = format_jam(meeting.waktu_selesai)
    durasi = hitung_durasi(meeting.waktu_mulai, meeting.waktu_selesai)
    pimpinan_text = f"{pimpinan.nama} ({pimpinan.jabatan})" if pimpinan else "pimpinan rapat"

    parts = [
        f"Pada hari {tanggal_lengkap}",
    ]
    if meeting.waktu_mulai:
        parts[0] += f", pukul {jam_mulai}"
    text = (
        f"{parts[0]}, telah dilaksanakan {meeting.judul_rapat}. "
        f"Rapat ini dipimpin langsung oleh {pimpinan_text}."
    )
    if durasi != "-":
        text += f" Kegiatan rapat berlangsung selama {durasi}"
        if meeting.waktu_selesai:
            text += f" dan berakhir pada pukul {jam_selesai}"
        text += "."
    doc.add_paragraph(text)


def _pembahasan_section(doc: Document, ringkasan, keputusan):
    doc.add_heading("Pembahasan Rapat", level=2)
    if ringkasan:
        for i, point in enumerate(ringkasan, start=1):
            doc.add_paragraph(f"{i}. {point}")
    else:
        doc.add_paragraph("Tidak ada poin pembahasan yang tercatat.")

    doc.add_heading("Keputusan Rapat", level=2)
    if keputusan:
        for point in keputusan:
            doc.add_paragraph(point, style="List Bullet")
    else:
        doc.add_paragraph("Tidak ada keputusan yang tercatat.")


def _tindak_lanjut_table(doc: Document, tindak_lanjut):
    doc.add_heading("Tindak Lanjut", level=2)
    if not tindak_lanjut:
        doc.add_paragraph("Tidak ada tindak lanjut yang teridentifikasi.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Action Item", "Penanggung Jawab", "Deadline", "Status"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    for item in tindak_lanjut:
        r = table.add_row().cells
        r[0].text = item["deskripsi"]
        r[1].text = item.get("penanggung_jawab") or "-"
        r[2].text = item.get("deadline") or "-"
        r[3].text = item.get("status") or "Pending"


def _tanda_tangan(doc: Document, meeting, kepala, notulis):
    doc.add_paragraph("")
    p = doc.add_paragraph(f"{(meeting.unit_kerja or settings.UNIT_KERJA_DEFAULT).split(' ')[-1]}, "
                           f"{format_tanggal_singkat(meeting.tanggal)}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    left_lines = ["Mengetahui,", f"Kepala {meeting.unit_kerja or settings.UNIT_KERJA_DEFAULT}", "", ""]
    right_lines = ["", "Notulis:", "", ""]

    for i, line in enumerate(left_lines):
        table.rows[i].cells[0].paragraphs[0].add_run(line)
    for i, line in enumerate(right_lines):
        table.rows[i].cells[1].paragraphs[0].add_run(line)

    # ruang tanda tangan
    for _ in range(2):
        table.add_row()

    nama_row = table.add_row().cells
    kepala_run = nama_row[0].paragraphs[0].add_run(kepala.nama if kepala else "-")
    kepala_run.bold = True
    kepala_run.underline = True
    notulis_run = nama_row[1].paragraphs[0].add_run(notulis.nama if notulis else "-")
    notulis_run.bold = True
    notulis_run.underline = True


def _dokumentasi_section(doc: Document, meeting, dokumentasi_files):
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f"DOKUMENTASI KEGIATAN {meeting.judul_rapat.upper()}")
    run.bold = True
    run.font.size = Pt(13)

    if not dokumentasi_files:
        p = doc.add_paragraph("Tidak ada bukti dokumentasi yang diunggah untuk kegiatan ini.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    # tampilkan foto 2 kolom per baris menggunakan tabel tanpa border
    photos = list(dokumentasi_files)
    rows_needed = (len(photos) + 1) // 2
    table = doc.add_table(rows=rows_needed, cols=2)
    idx = 0
    for r in range(rows_needed):
        for c in range(2):
            if idx >= len(photos):
                break
            cell = table.rows[r].cells[c]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            try:
                run.add_picture(str(photos[idx]["path"]), width=Cm(7))
            except Exception:
                para.add_run("[Gagal memuat gambar]")
            if photos[idx].get("keterangan"):
                cap = cell.add_paragraph(photos[idx]["keterangan"])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            idx += 1


def build_notulensi_docx(meeting, transcript_text, ringkasan, keputusan, tindak_lanjut,
                          peserta_list, pimpinan, notulis, kepala, dokumentasi_files, output_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_letterhead(doc)

    title = doc.add_heading(f"NOTULA KEGIATAN {meeting.judul_rapat.upper()}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _info_table(doc, meeting)
    _peserta_table(doc, peserta_list)
    doc.add_paragraph("")
    _pendahuluan_paragraph(doc, meeting, pimpinan)
    _pembahasan_section(doc, ringkasan, keputusan)
    _tindak_lanjut_table(doc, tindak_lanjut)
    _tanda_tangan(doc, meeting, kepala, notulis)
    _dokumentasi_section(doc, meeting, dokumentasi_files)

    if transcript_text:
        doc.add_page_break()
        doc.add_heading("Lampiran: Transkripsi Rekaman Rapat", level=2)
        doc.add_paragraph(transcript_text)

    doc.save(output_path)
    _fix_zoom_schema_quirk(output_path)
    return output_path


def _fix_zoom_schema_quirk(docx_path: Path):
    """python-docx's bundled default template omits w:percent on <w:zoom>,
    which is technically required by the OOXML schema (cosmetic issue only —
    Word/LibreOffice open the file fine either way, but this keeps the output
    fully schema-valid)."""
    import os
    import zipfile
    import shutil
    import tempfile

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)  # hanya perlu path-nya; fd yang tertinggal terbuka mengunci
                       # file di Windows dan membuat shutil.move() di bawah gagal
                       # dengan WinError 32
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml" and b'w:zoom' in data and b'w:percent' not in data:
                data = data.replace(b'<w:zoom w:val="bestFit"/>', b'<w:zoom w:val="bestFit" w:percent="100"/>')
            zout.writestr(item, data)
    shutil.move(tmp_path, docx_path)
